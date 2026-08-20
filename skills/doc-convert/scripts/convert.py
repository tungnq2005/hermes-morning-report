#!/usr/bin/env python3
"""Convert documents/presentations between formats.

Google Workspace is the renderer of record. The .pptx/.docx this skill builds is an
intermediate: it is imported into Google Slides / Google Docs, and whatever file the
user wants is exported back out of Google. python-pptx and LibreOffice each draw a
document their own way -- that is why a deck that looked right on Linux came out wrong
in PowerPoint for Mac -- whereas a Google file renders identically on macOS, Windows,
iPad and the browser, and its exports carry Google's rendering with them.

Without an authorized Google token the local file is delivered instead, flagged with a
`google_unauthorized:rendered_locally` warning; `gslides`/`gdoc` then fail outright.

Usage:
  python3 convert.py --input <path-or-url> --to gslides|gdoc|pptx|docx|pdf|md \
      [--title "..."] [--subtitle "..."] [--image <path>]... [--min-slides 5] \
      [--no-google] [--outdir DIR]

Prints a JSON manifest to stdout. Outputs default to
skills/doc-convert/state/output-history/YYYY-MM-DD/<run-id>/.
"""
from __future__ import annotations

import argparse
import datetime
import json
import os
import secrets
import shutil
import subprocess
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import doc_io
import image_search
from doc_io import DocConvertError

try:
    import google_io
except Exception:  # google libs optional; targets that need them fail with a clear message
    google_io = None

SKILL_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
# gslides/gdoc deliver a Google link (plus a PDF copy); pptx/docx/pdf deliver a file
# that Google exported, so every target but `md` goes through Google when authorized.
TARGETS = ("pptx", "docx", "pdf", "md", "gdoc", "gslides")

# python-docx's default template leaves body text on the theme's minor font, which is
# Cambria. Cambria ships with Office but not with Linux, so the `pdf` target -- which
# renders through headless LibreOffice -- substituted a serif face that has no
# precomposed Vietnamese glyphs. Diacritics were then drawn as separate combining
# glyphs with their own advance width: "so voi" rendered as "so vo i" with the tone
# mark beside the letter instead of above it.
#
# Calibri is the safe pin: Word and PowerPoint have it, and LibreOffice maps it to the
# metric-compatible Carlito, which does carry precomposed Vietnamese (U+1EA1 etc.).
# The pptx target already renders correctly because its theme pins both the major and
# the minor font to Calibri.
DOCX_FONT = "Calibri"
# python-docx's numbering.xml draws bullets as U+F0B7 -- a private-use codepoint that
# only exists in the Symbol font. Linux maps Symbol to OpenSymbol, which has nothing at
# F0B7, so every bullet rendered as a tofu box. U+2022 is the real BULLET codepoint and
# lives in the body font.
SYMBOL_BULLET = "\uf0b7"   # private-use codepoint, Symbol font only
UNICODE_BULLET = "\u2022"  # BULLET, present in Calibri/Carlito


def new_run_dir(outdir: str | None) -> str:
    if outdir:
        os.makedirs(outdir, exist_ok=True)
        return outdir
    now = datetime.datetime.now()
    run = os.path.join(SKILL_DIR, "state", "output-history",
                       now.strftime("%Y-%m-%d"), f"{now.strftime('%H%M%S')}-{secrets.token_hex(4)}")
    os.makedirs(run, exist_ok=True)
    return run


def _pin_font(style, name: str) -> None:
    """Pin a style to a literal font, overriding the template's theme reference.

    Setting only ``style.font.name`` leaves the ``w:asciiTheme``/``w:hAnsiTheme``
    attributes in place, and both Word and LibreOffice prefer the theme reference
    over the literal one. The theme attributes have to go.
    """
    from docx.oxml.ns import qn

    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn("w:" + attr), name)
    for attr in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme"):
        theme_attr = qn("w:" + attr)
        if theme_attr in rfonts.attrib:
            del rfonts.attrib[theme_attr]


def _fix_bullet_glyph(document, name: str) -> None:
    """Swap the Symbol-font private-use bullet for a real U+2022 in the body font."""
    from docx.oxml.ns import qn

    try:
        numbering = document.part.numbering_part.element
    except (AttributeError, KeyError, NotImplementedError):
        return
    for lvl_text in numbering.iter(qn("w:lvlText")):
        if lvl_text.get(qn("w:val")) == SYMBOL_BULLET:
            lvl_text.set(qn("w:val"), UNICODE_BULLET)
    for rfonts in numbering.iter(qn("w:rFonts")):
        if rfonts.get(qn("w:ascii")) == "Symbol":
            for attr in ("ascii", "hAnsi", "cs"):
                rfonts.set(qn("w:" + attr), name)


def _docx_section_images(doc: dict, sections: list[dict], images: list) -> tuple[str | None, dict[int, str]]:
    """Map per-section pictures onto the headings they belong under.

    `images` is indexed by section the way build_pptx consumes it, but build_docx
    walks blocks (see below), so the translation section index -> heading ordinal
    happens here. A document that opens with prose before its first heading has an
    untitled section 0; its picture goes directly under the title.
    """
    if not images or not sections:
        return None, {}
    has_preamble = not sections[0].get("title")
    preamble = images[0] if has_preamble and images[0] else None
    offset = 1 if has_preamble else 0
    plan: dict[int, str] = {}
    for ordinal in range(len(sections) - offset):
        index = ordinal + offset
        if index < len(images) and images[index]:
            plan[ordinal] = images[index]
    return preamble, plan


def _reencode_for_docx(path: str) -> str:
    """Rewrite an image into a header python-docx can parse.

    python-docx only understands JFIF (APP0) and Exif (APP1) JPEGs. Flickr -- where
    most Openverse hits live -- serves plenty of files whose first marker is APP13
    (a Photoshop resource block); python-pptx takes them, python-docx raises
    UnrecognizedImageError, and the picture silently vanished from documents.
    Pillow ships as a python-pptx dependency, so re-encoding costs no new install.
    """
    from PIL import Image

    out = os.path.splitext(path)[0] + "-docx.jpg"
    with Image.open(path) as img:
        img.convert("RGB").save(out, "JPEG", quality=88)
    return out


def _add_docx_picture(d, path: str, rejected: list[str]) -> bool:
    """Place one centred picture. A bad file loses its picture, not the document."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    for candidate in (path, None):
        if candidate is None:
            try:
                candidate = _reencode_for_docx(path)
            except Exception:  # noqa: BLE001 - not an image Pillow can read either
                break
        try:
            d.add_picture(candidate, width=Inches(5.5))
        except Exception:  # noqa: BLE001 - unreadable/unsupported image header
            continue
        d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    rejected.append(os.path.basename(path))
    return False


def _docx_credits(d, doc: dict, credits: list[dict]) -> None:
    """CC-BY images must name their creator, in Word as much as on a slide."""
    import build_pptx
    from docx.shared import Pt, RGBColor

    words = build_pptx.STRINGS["en" if build_pptx._is_english(doc) else "vi"]
    heading = d.add_heading(words["credits"], level=2)
    for run in heading.runs:
        run.font.name = DOCX_FONT
    for credit in credits:
        title = (credit.get("title") or credit.get("query") or "Untitled")[:60]
        creator = credit.get("creator") or "Unknown"
        licence = credit.get("license") or "CC"
        p = d.add_paragraph(f"{title} — {creator} ({licence}) · Openverse")
        for run in p.runs:
            run.font.name = DOCX_FONT
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def build_docx(doc: dict, out_path: str, *, sections: list[dict] | None = None,
               images: list | None = None, credits: list[dict] | None = None) -> dict:
    import docx
    from docx.shared import Pt, RGBColor

    d = docx.Document()
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3",
                       "Heading 4", "List Bullet"):
        try:
            _pin_font(d.styles[style_name], DOCX_FONT)
        except KeyError:
            pass
    _fix_bullet_glyph(d, DOCX_FONT)

    h = d.add_heading(doc["title"], level=0)
    for run in h.runs:
        run.font.name = DOCX_FONT
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    rejected: list[str] = []
    used = 0
    preamble_image, image_plan = _docx_section_images(doc, sections or [], images or [])
    if preamble_image and _add_docx_picture(d, preamble_image, rejected):
        used += 1

    # Walk the blocks rather than the slide outline. `outline_sections` flattens every
    # paragraph into sentence-sized `items` because that is what slides want; reusing it
    # here turned every line of prose into a bullet and split long paragraphs across
    # several of them, so the Word file held no body text at all.
    heading_ordinal = 0
    for block in doc["blocks"]:
        if block["kind"] == "heading":
            heading = d.add_heading(block["text"], level=min(block.get("level", 1), 4))
            for run in heading.runs:
                run.font.name = DOCX_FONT
            picture = image_plan.get(heading_ordinal)
            heading_ordinal += 1
            if picture and _add_docx_picture(d, picture, rejected):
                used += 1
            continue
        style = "List Bullet" if block["kind"] == "bullet" else None
        p = d.add_paragraph(block["text"], style=style) if style else d.add_paragraph(block["text"])
        for run in p.runs:
            run.font.name = DOCX_FONT
            run.font.size = Pt(11)

    if used and credits:
        _docx_credits(d, doc, credits)
    d.save(out_path)
    return {"images_used": used, "images_rejected": rejected}


def resolve_images(args, doc: dict, sections: list[dict], run_dir: str,
                   manifest: dict, build_pptx) -> tuple[list, list]:
    """Decide which pictures the deck gets. Never returns an irrelevant one.

    Explicit --image wins. Otherwise we search Openverse, but only with English
    queries: the agent supplies --image-query per section, or the section titles
    themselves serve when the document is already English. A Vietnamese title is
    left unsearched -- Openverse answers it with fishing boats -- and the slide
    simply goes without a picture.
    """
    if args.image:
        return list(args.image), []
    if args.no_auto_images:
        return [], []

    queries = list(args.image_query or [])
    if not queries:
        if not build_pptx._is_english(doc):
            manifest["warnings"].append("image_search_needs_english_query")
            return [], []
        queries = [sec["title"] or doc["title"] for sec in sections]

    paths, credits, warnings = image_search.fetch(queries, os.path.join(run_dir, "images"))
    manifest["warnings"].extend(warnings)
    if credits:
        manifest["image_credits"] = credits
    return paths, credits


def to_pdf(src_path: str, run_dir: str) -> str:
    cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", run_dir, src_path]
    proc = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
    expected = os.path.join(run_dir, os.path.splitext(os.path.basename(src_path))[0] + ".pdf")
    if proc.returncode != 0 or not os.path.exists(expected):
        raise DocConvertError(f"LibreOffice PDF conversion failed: {(proc.stderr or proc.stdout).strip()[:300]}")
    return expected


def local_artifact(args, doc: dict, sections: list[dict], src: str, src_ext: str,
                   run_dir: str, build_dir: str, base: str, manifest: dict, want: str) -> str:
    """Build (or reuse) the local .pptx/.docx that Google will import.

    An input that is already in the wanted shape is handed to Google untouched: the
    user asking to put *their* deck on Google Slides wants that deck, not our
    regenerated one. `--rebuild` is for when they want ours (re-laid out, with imagery).
    """
    if src_ext == "." + want and not args.rebuild:
        return src
    os.makedirs(build_dir, exist_ok=True)
    out_path = os.path.join(build_dir, base + "." + want)
    if want == "pptx":
        import build_pptx

        images, credits = resolve_images(args, doc, sections, run_dir, manifest, build_pptx)
        stats = build_pptx.build(doc, sections, out_path, min_slides=args.min_slides,
                                 images=images, subtitle=args.subtitle, credits=credits)
        # Stat cards and card grids have no room for a photo, so a number-heavy
        # document (a morning report, say) came out with imagery fetched and none of
        # it shown. Rebuild once with the first picture on the cover instead of
        # shipping a deck that quietly dropped every image it downloaded. Decks that
        # already place pictures are left exactly as they were.
        cover = next((p for p in images if p), None)
        if cover and not stats.get("images_used"):
            stats = build_pptx.build(doc, sections, out_path, min_slides=args.min_slides,
                                     images=images, subtitle=args.subtitle, credits=credits,
                                     cover_image=cover)
        for name in stats.pop("images_rejected", []):
            manifest["warnings"].append(f"image_embed_failed:{name}")
        manifest.update(stats)
    else:
        # Documents take pictures only when the caller asked for them. A deck is
        # expected to be illustrated; a Word file usually is not, and nobody wants
        # stock photography appearing inside a contract they only asked to convert.
        images: list = []
        credits: list[dict] = []
        if (args.image or args.image_query) and not args.no_auto_images:
            import build_pptx

            images, credits = resolve_images(args, doc, sections, run_dir, manifest, build_pptx)
        stats = build_docx(doc, out_path, sections=sections, images=images, credits=credits)
        for name in stats.pop("images_rejected", []):
            manifest["warnings"].append(f"image_embed_failed:{name}")
        manifest.update(stats)
    return out_path


def local_fallback(local_path: str, run_dir: str, target: str) -> str | None:
    """What ships when Google cannot: our own file, or LibreOffice's PDF.

    For `gslides`/`gdoc` there is nothing to ship -- the link is the deliverable and
    the PDF was only the offline copy -- so this returns None and the caller's warning
    stands on its own.
    """
    if target == "pdf":
        return to_pdf(local_path, run_dir)
    if target in ("pptx", "docx"):
        dest = os.path.join(run_dir, os.path.basename(local_path))
        if os.path.abspath(local_path) != os.path.abspath(dest):
            shutil.copy2(local_path, dest)
        return dest
    return None


def finish(manifest: dict, run_dir: str) -> int:
    with open(os.path.join(run_dir, "manifest.json"), "w", encoding="utf-8") as fh:
        json.dump(manifest, fh, ensure_ascii=False, indent=2)
    print(json.dumps(manifest, ensure_ascii=False, indent=2))
    return 0 if manifest.get("success") else 1


def check_google_deck(pres_id: str) -> dict:
    """Read the imported deck back out of Google and re-run the layout invariants.

    Drive's importer re-maps placeholders and re-wraps text with its own font
    substitute, so a clean local .pptx is no evidence about the deck Google serves.
    Returns status pass/fail/unchecked -- `unchecked` is not a pass.
    """
    try:
        import validate_output

        data = google_io.inspect_presentation(pres_id)
        problems = validate_output.check_google_slides(data, validate_output._font_path())
        return {"status": "fail" if problems else "pass",
                "slides": len(data.get("slides") or []), "problems": problems}
    except Exception as err:  # noqa: BLE001 - a failed probe must not lose the deck
        return {"status": "unchecked", "error": f"{type(err).__name__}: {err}"}


def main() -> int:
    ap = argparse.ArgumentParser(description="Convert documents between formats")
    ap.add_argument("--input", required=True, help="Local file path or public URL / Google link")
    ap.add_argument("--to", required=True, choices=TARGETS)
    ap.add_argument("--title", default="", help="Override document title")
    ap.add_argument("--subtitle", default="", help="Subtitle for the pptx title slide")
    ap.add_argument("--image", action="append", default=[], help="Image file to place on slides (repeatable)")
    ap.add_argument("--image-query", action="append", default=[],
                    help="English search phrase for one section's slide image, in section order "
                         "(repeatable). Openverse returns nothing useful for Vietnamese queries.")
    ap.add_argument("--no-auto-images", action="store_true",
                    help="Never search Openverse; leave slides without pictures.")
    ap.add_argument("--min-slides", type=int, default=5)
    ap.add_argument("--rebuild", action="store_true",
                    help="Re-lay out an input that is already in the target shape "
                         "(a .pptx asked for as slides is otherwise uploaded as-is).")
    ap.add_argument("--no-google", action="store_true",
                    help="Render locally even when Google is authorized (offline/debug). "
                         "The resulting file is python-pptx/LibreOffice output, which is "
                         "what renders inconsistently in PowerPoint for Mac.")
    ap.add_argument("--outdir", default=None)
    args = ap.parse_args()

    manifest: dict = {"success": False, "input": args.input, "target": args.to, "warnings": []}
    try:
        run_dir = new_run_dir(args.outdir)
        manifest["run_dir"] = run_dir

        src = args.input
        if doc_io.is_url(src):
            # Private Google file: use authorized API if a token exists; else public download.
            # Only take the authorized route when the token can actually read files the
            # bot did not create; a minimal-scope deployment falls back to the public
            # download, which still works for a shared link.
            if (google_io and google_io.is_google_url(src) and google_io.has_token()
                    and google_io.can_read_private_files()):
                src = google_io.download_private(src, os.path.join(run_dir, "input"))
                manifest["downloaded_to"] = src
                manifest["source_access"] = "google-authorized"
            else:
                src = doc_io.download(src, os.path.join(run_dir, "input"))
                manifest["downloaded_to"] = src
        elif not os.path.exists(src):
            raise DocConvertError(f"Input file not found: {src}")
        else:
            # keep a copy so the run dir is self-contained and under the workspace
            copied = os.path.join(run_dir, "input", os.path.basename(src))
            os.makedirs(os.path.dirname(copied), exist_ok=True)
            shutil.copy2(src, copied)
            src = copied

        src_ext = doc_io.detect_ext(src)
        manifest["detected_type"] = src_ext
        if src_ext == "." + args.to:
            manifest["warnings"].append("input and target formats are the same")

        base = os.path.splitext(os.path.basename(src))[0]
        build_dir = os.path.join(run_dir, "build")

        google_ok = bool(google_io) and google_io.has_token() and not args.no_google
        manifest["google_available"] = google_ok

        # Markdown never goes near Google: it is plain text, identical everywhere.
        if args.to == "md":
            doc = doc_io.extract(src)
            if args.title:
                doc["title"] = args.title
            out_path = os.path.join(run_dir, base + ".md")
            with open(out_path, "w", encoding="utf-8") as fh:
                fh.write(doc_io.to_markdown(doc))
            manifest.update({"success": True, "output": out_path, "title": doc["title"],
                             "render_engine": "local"})
            return finish(manifest, run_dir)

        if args.to in ("gslides", "gdoc") and not google_ok:
            if not google_io:
                raise DocConvertError("Thiếu Google API libs. Cài: pip3 install google-api-python-client google-auth-oauthlib")
            if args.no_google:
                raise DocConvertError("--no-google loại trừ chính target đang yêu cầu (gslides/gdoc).")
            raise DocConvertError(
                "Chưa kết nối Google. Kết nối ngay trong chat: skill_view(name=\"guided-setup\") "
                "-> Connect Google. (Đường terminal cho người vận hành: authorize_google.py)")

        # Which Office shape Google has to import: a deck for slide targets, a document
        # for the rest. `pdf` follows the input so a deck stays a deck.
        want = "pptx" if args.to in ("pptx", "gslides") else "docx"
        if args.to == "pdf" and src_ext == ".pptx":
            want = "pptx"

        # An office file that only needs a PDF keeps its own layout: never re-extract it.
        passthrough = args.to == "pdf" and src_ext in (".docx", ".pptx")
        if passthrough:
            local_path, doc = src, {"title": args.title or base}
        else:
            doc = doc_io.extract(src)
            if args.title:
                doc["title"] = args.title
            sections = doc_io.outline_sections(doc)
            manifest["sections"] = len(sections)
            local_path = local_artifact(args, doc, sections, src, src_ext, run_dir, build_dir,
                                        base, manifest, want)
        manifest["local_build"] = local_path

        if google_ok:
            # Google renders the deliverable. python-pptx/LibreOffice output is only ever
            # an intermediate now: that output is what rendered wrong in PowerPoint for
            # Mac, and a Google-exported file carries Google's own rendering instead.
            kind = "gslides" if want == "pptx" else "gdoc"
            imported = google_io.import_local(local_path, kind, title=doc.get("title", base))
            manifest.update({"google_url": imported["url"], "google_id": imported["id"],
                             "google_kind": kind, "render_engine": "google"})
            if kind == "gslides":
                manifest["google_check"] = check_google_deck(imported["id"])
                slides = manifest["google_check"].get("slides")
                if slides:
                    manifest["slides"] = slides

            # gslides/gdoc are delivered as a link; the PDF rides along as the offline copy.
            fmt = "pdf" if args.to in ("gslides", "gdoc", "pdf") else args.to
            dest = os.path.join(run_dir, base + google_io.EXPORT_MIMES[fmt][1])
            try:
                out_path = google_io.export_to(imported["id"], fmt, dest)
            except google_io.GoogleExportError as err:
                manifest["warnings"].append(f"google_export_failed:{err}")
                out_path = local_fallback(local_path, run_dir, args.to)
        else:
            manifest["warnings"].append(
                "google_disabled:rendered_locally" if args.no_google
                else "google_unauthorized:rendered_locally")
            manifest["render_engine"] = "local"
            out_path = local_fallback(local_path, run_dir, args.to)

        manifest.update({"success": True, "title": doc.get("title", base)})
        if out_path:
            manifest["output"] = out_path
    except DocConvertError as err:
        manifest["error"] = str(err)
    except Exception as err:  # unexpected - keep JSON contract for the agent
        manifest["error"] = f"{type(err).__name__}: {err}"

    return finish(manifest, manifest.get("run_dir", "."))


if __name__ == "__main__":
    sys.exit(main())
